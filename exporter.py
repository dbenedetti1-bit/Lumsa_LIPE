from __future__ import annotations

from pathlib import Path

from docx import Document


def save_markdown(md_path: str | Path, content: str) -> Path:
    md_path = Path(md_path)
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(content.strip() + "\n", encoding="utf-8")
    return md_path


def save_docx(docx_path: str | Path, markdown_like_text: str, *, title: str = "Trascrizione") -> Path:
    """
    Esporta un testo "markdown-like" in DOCX con parsing minimale di headings/bullet.
    (Non è un renderer Markdown completo: è volutamente semplice e robusto.)
    """
    docx_path = Path(docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=0)

    for raw in markdown_like_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        if line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(line)

    doc.save(str(docx_path))
    return docx_path

